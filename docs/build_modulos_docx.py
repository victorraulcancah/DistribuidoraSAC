from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "modulos-y-submodulos-sistemas.md"
OUTPUT = ROOT / "modulos-y-submodulos-sistemas.docx"


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.font.bold = True


def set_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_summary_table(doc):
    rows = [
        ("ERP", "Negocio, finanzas, compras, ventas y administracion"),
        ("WMS", "Almacen, inventario operativo y ejecucion logistica interna"),
        ("TMS", "Transporte, rutas, flota, entregas y distribucion"),
        ("DMS", "Documentos, versiones, aprobaciones, firma y archivo"),
        ("RRHH", "Empleados, asistencia, nomina, desempeno y portales"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.4)
    table.columns[1].width = Inches(5.0)
    set_cell(table.rows[0].cells[0], "Sistema", bold=True)
    set_cell(table.rows[0].cells[1], "Responsabilidad principal", bold=True)
    for system, responsibility in rows:
        cells = table.add_row().cells
        set_cell(cells[0], system, bold=True)
        set_cell(cells[1], responsibility)


def add_markdown_content(doc, markdown):
    in_summary_table = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "## Resumen de responsabilidades":
            doc.add_heading("Resumen de responsabilidades", level=1)
            add_summary_table(doc)
            in_summary_table = True
            continue
        if in_summary_table and line.startswith("|"):
            continue
        if in_summary_table and not line.startswith("|"):
            in_summary_table = False

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            heading = line[3:]
            if heading.startswith(("2. ", "3. ", "4. ", "5. ")):
                doc.add_page_break()
            doc.add_heading(heading, level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def add_footer(doc):
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Modulos y submodulos de sistemas")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Modulos y submodulos de los sistemas")
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("ERP, WMS, TMS, DMS y RRHH")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")

    add_markdown_content(doc, SOURCE.read_text(encoding="utf-8"))
    add_footer(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
